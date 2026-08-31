"""Extract plain text from office / spreadsheet / text uploads."""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger("doqseal.office")


OFFICE_EXTENSIONS = {
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
    ".txt",
    ".md",
}


def is_office_or_text(mime_type: str, filename: str = "") -> bool:
    mime = (mime_type or "").lower()
    ext = Path(filename or "").suffix.lower()
    if ext in OFFICE_EXTENSIONS:
        return True
    markers = (
        "wordprocessingml",
        "msword",
        "spreadsheetml",
        "ms-excel",
        "csv",
        "text/plain",
        "text/markdown",
        "officedocument",
    )
    return any(m in mime for m in markers)


def extract_office_text(
    file_bytes: bytes,
    mime_type: str,
    filename: str = "",
    *,
    max_chars: int = 60_000,
) -> str:
    mime = (mime_type or "").lower()
    ext = Path(filename or "").suffix.lower()

    try:
        if ext == ".docx" or "wordprocessingml" in mime or mime.endswith("msword"):
            return _docx_text(file_bytes, max_chars)
        if ext in {".xlsx", ".xls"} or "spreadsheetml" in mime or "ms-excel" in mime:
            return _xlsx_text(file_bytes, max_chars)
        if ext == ".csv" or "csv" in mime:
            return _csv_text(file_bytes, max_chars)
        if ext in {".txt", ".md"} or mime.startswith("text/"):
            return file_bytes.decode("utf-8", errors="ignore")[:max_chars]
    except Exception as err:
        logger.warning("Office text extract failed (%s): %s", filename or mime, err)
        raise

    raise ValueError(f"Unsupported office/text type: {mime or ext or 'unknown'}")


def _docx_text(file_bytes: bytes, max_chars: int) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("DOCX contained no extractable text")
    return text[:max_chars]


def _xlsx_text(file_bytes: bytes, max_chars: int) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in wb.worksheets:
            parts.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
                if sum(len(p) for p in parts) >= max_chars:
                    break
            if sum(len(p) for p in parts) >= max_chars:
                break
    finally:
        wb.close()
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Spreadsheet contained no extractable cells")
    return text[:max_chars]


def _csv_text(file_bytes: bytes, max_chars: int) -> str:
    raw = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(raw))
    parts: list[str] = []
    for row in reader:
        cells = [c.strip() for c in row if c and c.strip()]
        if cells:
            parts.append(" | ".join(cells))
        if sum(len(p) for p in parts) >= max_chars:
            break
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("CSV contained no extractable rows")
    return text[:max_chars]
